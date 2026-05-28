"""Generate `P02_REPORT.docx` from the official template + our content.

We open the provided P02 template so that page setup, margins and default
fonts are inherited from the file the teacher expects. We then wipe the
placeholder body and write the report's sections natively (Heading 1 +
body paragraphs + occasional bullet list / table / monospace block).

Run with:  .\.venv\Scripts\python.exe build_p02_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


TEMPLATE = "P02 Report Template - Data ProcVisualization.docx"
OUT = "P02_REPORT.docx"


def clear_body(doc):
    """Remove every paragraph and table currently in the document body."""
    body = doc.element.body
    for child in list(body):
        # Keep section properties (page size, margins). Those live in <w:sectPr>.
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_para(doc, text, *, style=None, bold=False, italic=False, mono=False,
             align=None, size=None, color=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if bold:   run.bold = True
    if italic: run.italic = True
    if mono:
        run.font.name = "Consolas"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run("•  ")
        # Allow simple inline bold via **...** at the start "key — value"
        if "**" in item:
            parts = item.split("**")
            for i, chunk in enumerate(parts):
                run = p.add_run(chunk)
                run.bold = (i % 2 == 1)
        else:
            p.add_run(item)


def add_code_block(doc, lines):
    """Render a multi-line code block (no syntax highlighting — just monospace)."""
    for line in lines:
        add_para(doc, line, mono=True, size=10, color="333333")


def add_simple_table(doc, header, rows):
    tbl = doc.add_table(rows=1, cols=len(header))
    # Try a couple of common built-in styles; fall back to plain table.
    for candidate in ("Table Grid", "Light List Accent 1", "Light Grid"):
        try:
            tbl.style = candidate
            break
        except KeyError:
            continue
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return tbl


# =============================================================================
# Build
# =============================================================================

assert Path(TEMPLATE).exists(), f"missing template: {TEMPLATE}"
doc = Document(TEMPLATE)
clear_body(doc)

# ---------------- Title block ----------------
add_para(doc, "Data Processing and Visualization (P02)",
         bold=True, size=22, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_para(doc, "Decision Support Systems, 2025–26",
         italic=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_para(doc, "Team ###: First name (number), First name (number), First name (number)",
         italic=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_para(doc, "")

# ---------------- 1. Introduction ----------------
add_para(doc, "1. Introduction", style="Heading 1")

add_para(doc,
    "P01 left us with a working data mart on Supabase: a Medallion pipeline that "
    "lands raw Wide World Importers data in bronze, cleans it into silver, and "
    "finally serves a star schema in gold with seven dimensions and two fact "
    "tables. What was missing was the layer that turns those tables into "
    "something a person can actually look at and reason about. That is what P02 "
    "is for."
)
add_para(doc,
    "We built an interactive dashboard with Streamlit and Plotly, reading "
    "directly from the gold schema. The choice came down to staying inside the "
    "Python stack we already had — same SQLAlchemy connection, same .env, same "
    "repository — instead of jumping to Power BI and re-importing the model. "
    "It also keeps the delivery simple: one command spins the app up locally, "
    "and a few clicks publish it to Streamlit Community Cloud."
)
add_para(doc,
    "The dashboard is built around six business questions that the dimensional "
    "model is well-suited for. Which products and customer categories drive "
    "revenue and profit. How revenue is distributed across territories. Who the "
    "top customers are. Which salespersons perform best. And where "
    "outstanding balance is sitting at any moment."
)
add_para(doc,
    "The intended audience is split: sales managers care about the revenue "
    "trend and the top-N views; account managers and marketing live on the "
    "Customers page; sales operations look at salesperson rankings; "
    "and finance / accounts receivable use the Finance page for "
    "outstanding balance analysis."
)

# ---------------- 2. Data acquisition and preparation ----------------
add_para(doc, "2. Data acquisition and preparation", style="Heading 1")

add_para(doc,
    "There is no new ETL in P02. The dashboard reads only from gold.* — the "
    "same tables P01 already produces. What did happen during setup, though, "
    "was that we wrote a validator (validate_db.py) to cross-check the warehouse "
    "before plugging the dashboard in, and it caught a problem that had slipped "
    "through."
)
add_para(doc,
    "The row counts came back inflated. bronze.invoices had 141 020 rows when "
    "the source has 70 510. gold.factsales had 913 060 instead of 228 265 — a "
    "clean 2× on the invoice header and 4× on the line items. Looking at "
    "bronze._load_control, both invoices and invoicelines showed n=2 for the "
    "incremental strategy: someone had run the incremental load twice. The "
    "watermark had not been bumped between the two runs, the second pass "
    "thought the same window was still new, and there was no de-duplication "
    "downstream. Each invoice ended up in bronze twice, and when silver joined "
    "invoice lines to invoice headers, the duplication amplified."
)
add_para(doc,
    "The cleanup is packaged in fix_dup.sql: truncate gold.factsales, "
    "gold.factinvoices, the silver fact-staging tables, bronze.invoices and "
    "bronze.invoicelines, and the matching rows in bronze._load_control. After "
    "that, re-running the incremental cells once (01_bronze.ipynb ## 9 and "
    "## 10), then 02_silver.ipynb and 03_gold.ipynb, brought everything back "
    "to the expected numbers. validate_db.py then reported zero failures across "
    "all eight checks: orphan FKs, unmapped FKs, duplicate current rows, NULL "
    "percentage on business keys, DimDate coverage, silver→gold reconciliation "
    "and the status of every entry in _load_control."
)
add_para(doc,
    "On the dashboard side, “preparation” really just means caching. We do not "
    "pre-aggregate anything in the database — the volume is small enough that "
    "Supabase answers each query in well under a second on the indexed gold "
    "tables. Each query function is wrapped with @st.cache_data(ttl=600), and "
    "the SQLAlchemy engine itself with @st.cache_resource, so a rerun triggered "
    "by a slider change only re-executes the queries whose filter arguments "
    "actually moved."
)
add_para(doc,
    "Credentials live in .env during local development and in st.secrets when "
    "deployed on Streamlit Community Cloud. A small helper in dashboard/db.py "
    "(_secret()) tries st.secrets first and falls back to the environment, so "
    "the same code runs unchanged in both places."
)

# ---------------- 3. Data modelling and processing ----------------
add_para(doc, "3. Data modelling and processing", style="Heading 1")

add_para(doc,
    "We kept the model exactly as P01 left it: six SCD-Type-2 dimensions plus "
    "a static DimDate, and two fact tables (FactSales at the invoice-line "
    "grain, FactInvoices at the invoice-header grain). No calculated columns "
    "were added to gold. Every measure shown in the dashboard is computed at "
    "query time. That keeps the warehouse immutable and lets the sidebar "
    "filters compose naturally with whatever aggregation each chart needs."
)
add_para(doc, "The measures used across the app:")
add_bullets(doc, [
    "**Revenue** — SUM(factsales.extendedprice)",
    "**Profit** — SUM(factsales.lineprofit)",
    "**Gross margin %** — profit / revenue × 100",
    "**Quantity sold** — SUM(factsales.quantity)",
    "**Invoices count** — COUNT(*) over factinvoices",
    "**Average invoice amount** — AVG(factinvoices.invoiceamount)",
    "**Outstanding balance** — SUM(factinvoices.outstandingbalance)",
    "**Outstanding ratio** — outstanding / invoiced × 100",
    "**Active customers and products sold** — COUNT(DISTINCT ...key) in factsales",
])
add_para(doc,
    "Two hierarchies show up across the charts. On the date axis we use "
    "Year → Month → Day (driven by DimDate). On the location axis we use "
    "Sales Territory → Country (and City is available in the dim if a future "
    "page wants to drill deeper). The customer side has Category → Customer; "
    "on the product side, Brand → Product. There is no product category in "
    "the source — only a customer category — so the treemap on the Customers "
    "page uses Customer Category × Product Brand as the closest analogue we "
    "could build."
)
add_para(doc,
    "All joins between facts and SCD-2 dimensions filter on date_to IS NULL "
    "so the dashboard always reflects the current attributes of each entity. "
    "We isolated the join pattern in two SQL fragments inside "
    "dashboard/queries.py (_BASE_SALES, _BASE_INVOICES); every chart on every "
    "page reuses one of them, which keeps the semantics consistent and the "
    "SQL readable."
)

# ---------------- 4. Data visualization ----------------
add_para(doc, "4. Data visualization", style="Heading 1")

add_para(doc,
    "The app is organised into a Home and four themed pages. The sidebar holds "
    "the global filters — year range, customer category and sales territory. "
    "Selections are persisted in st.session_state, so changing them on one "
    "page is felt on the others."
)
add_para(doc,
    "The Home is the executive summary. Eight KPI cards across two rows "
    "(revenue, profit, invoices, average invoice, total quantity, "
    "outstanding balance, active customers, products sold), plus the gross "
    "margin in a caption. Below the KPIs, a single time-series chart with "
    "monthly revenue and profit. It is the page someone opens to know whether "
    "things are roughly OK before drilling anywhere."
)
add_para(doc,
    "Sales Performance drills into where the revenue is coming from. The "
    "monthly revenue and profit line appears again, and underneath it a stacked "
    "area showing revenue by customer category over time. The lower half of "
    "the page is split: on the left, a Top-N product bar with a slider going "
    "from 5 to 50 and a detail table behind an expander; on the right, a "
    "horizontal bar of countries coloured by sales territory."
)
add_para(doc,
    "Customers shows a straightforward Top-N bar of customers ranked by "
    "revenue, coloured by category, so it is easy to spot who the biggest "
    "accounts are. A formatted detail table follows it. At the bottom, a "
    "treemap of customer category × product brand for the revenue-mix view."
)
add_para(doc,
    "Operations focuses on salesperson productivity. Two bar charts rank the "
    "team by revenue and by quantity sold, so managers can see whether the "
    "same people lead in both. A full detail table with revenue, profit, "
    "invoices and quantity rounds the page out."
)
add_para(doc,
    "Finance focuses on receivables. Three KPIs at the top (outstanding "
    "balance, total invoiced, outstanding ratio), a bar of outstanding by "
    "sales territory, and a scatter of outstanding vs. invoiced per customer "
    "coloured by category. A detail table is available behind an expander."
)
add_para(doc,
    "A few small things we did across the app to keep it clean: a single "
    "Plotly template (plotly_white) and one qualitative palette (Set2) shared "
    "by every chart; monetary values abbreviated as K and M in KPI cards but "
    "kept whole in tables and hover; and a small helper in dashboard/charts.py "
    "that draws a friendly “No data for the selected filters” message instead "
    "of letting a chart fail when filters return nothing."
)
add_para(doc,
    "Screenshots of each page are in the assets/ folder and referenced in "
    "the team submission."
)

# ---------------- 5. Conclusion ----------------
add_para(doc, "5. Conclusion", style="Heading 1")

add_para(doc,
    "The project finished with a focused decision-support layer on "
    "top of the P01 warehouse: around 15 charts across 5 pages, answering 6 "
    "business questions, with all the underlying numbers verified by a "
    "validator that returns zero failures end-to-end."
)
add_para(doc,
    "The thing that took the most time was not the visualisation. It was "
    "catching and fixing the duplicate-load bug we found before plugging the "
    "dashboard in. Without validate_db.py, the dashboard would have shown "
    "sales totals four times higher than they should be, invoice totals twice "
    "as high, and we probably would not have noticed until someone tried to "
    "reconcile a KPI against an external source. That validator is worth "
    "keeping around independently of the dashboard, and it is the artefact we "
    "are most pleased with."
)
add_para(doc,
    "Streamlit and Plotly delivered what we wanted. Going from “data mart is "
    "ready” to “interactive dashboard in the browser” took an afternoon once "
    "the queries were sketched out, and there is essentially nothing to "
    "maintain on the front-end side: no JavaScript, no asset bundling, no CSS. "
    "The cache decorators kept things responsive even when sliders move "
    "quickly, and the parameterised SQL means the same query function powers "
    "four or five charts depending on the filters it receives."
)
add_para(doc,
    "Things we would do differently or add later. Pre-aggregating monthly "
    "snapshots if the volume grew past a few million rows. Adding a "
    "forecasting tab with a simple time-series model. And finally addressing a "
    "credential hygiene issue carried over from the original P01 repository — "
    "an early commit by a teammate contained the .env file, which we worked "
    "around by starting a fresh repository (zprog10/project_DSS) for this "
    "delivery and deploying from there. Rotating the Supabase password is the "
    "right next step after submission and was left out of scope to keep this "
    "report focused."
)
add_para(doc,
    "The app is reachable on Streamlit Community Cloud at the URL the team "
    "paste in the cover page of the submission. To reproduce a local install, "
    "pip install -r requirements.txt followed by streamlit run app.py is "
    "enough, assuming a .env with the Supabase credentials."
)

# ---------------- Appendix ----------------
add_para(doc, "Appendix — How to run", style="Heading 1")
add_code_block(doc, [
    "# 1. Install",
    ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt",
    "",
    "# 2. Verify the warehouse (should end with \"Tudo OK\")",
    ".\\.venv\\Scripts\\python.exe validate_db.py",
    "",
    "# 3. Launch the dashboard",
    ".\\.venv\\Scripts\\streamlit.exe run app.py",
])
add_para(doc, "")
add_para(doc, "After a fresh load expect these numbers; the Home KPIs should match:")
add_simple_table(doc,
    header=["Table", "Expected rows"],
    rows=[
        ["bronze.invoices",     "70 510"],
        ["bronze.invoicelines", "228 265"],
        ["gold.factsales",      "228 265"],
        ["gold.factinvoices",   "70 510"],
    ],
)
add_para(doc, "Home → Total Invoices (no filters) should read 70 510.")

doc.save(OUT)
print(f"Wrote {OUT}")
