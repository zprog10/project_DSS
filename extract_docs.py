import pdfplumber
from docx import Document

print("=" * 80)
print("DSS91_Project01_Guide_ESI.pdf")
print("=" * 80)
try:
    with pdfplumber.open("DSS91_Project01_Guide_ESI.pdf") as pdf:
        for i, page in enumerate(pdf.pages, 1):
            print(f"\n--- PAGE {i} ---")
            text = page.extract_text()
            if text:
                print(text)
except Exception as e:
    print(f"Error reading PDF: {e}")

print("\n\n" + "=" * 80)
print("P01 Report Template - Data Mart.docx")
print("=" * 80)
try:
    doc = Document("P01 Report Template - Data Mart.docx")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    for table in doc.tables:
        print("\n--- TABLE ---")
        for row in table.rows:
            print(" | ".join(cell.text for cell in row.cells))
except Exception as e:
    print(f"Error reading DOCX: {e}")
